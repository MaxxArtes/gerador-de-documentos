# backend/main.py

from fastapi import FastAPI, Form
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fpdf import FPDF
from docx import Document
from bs4 import BeautifulSoup
import base64
import os

app = FastAPI()

# Servir o frontend
@app.get("/", response_class=HTMLResponse)
async def read_root():
    frontend_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "../frontend/index.html"))
    with open(frontend_path, "r", encoding="utf-8") as f:
        html_content = f.read()
    return HTMLResponse(content=html_content)

@app.post("/gerar_documento/")
async def gerar_documento(conteudo: str = Form(...), formato: str = Form(...)):
    # Limpa arquivos antigos
    for f in ["documento_gerado.pdf", "documento_gerado.docx", "documento_gerado.txt", "documento_gerado.md"]:
        if os.path.exists(f):
            os.remove(f)

    # Processa HTML recebido
    soup = BeautifulSoup(conteudo, 'html.parser')

    if formato == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)

        for elem in soup.descendants:
            if elem.name == "img":
                img_data = elem['src'].split(",")[-1]
                with open("temp_image.png", "wb") as img_file:
                    img_file.write(base64.b64decode(img_data))
                pdf.image("temp_image.png", w=100)
            elif elem.name is None and elem.string and elem.string.strip():
                pdf.multi_cell(0, 10, elem.string.strip())

        if os.path.exists("temp_image.png"):
            os.remove("temp_image.png")

        pdf.output("documento_gerado.pdf")
        return FileResponse("documento_gerado.pdf", media_type="application/pdf", filename="documento_gerado.pdf")

    elif formato == "docx":
        doc = Document()
        for elem in soup.descendants:
            if elem.name == "img":
                img_data = elem['src'].split(",")[-1]
                with open("temp_image.png", "wb") as img_file:
                    img_file.write(base64.b64decode(img_data))
                doc.add_picture("temp_image.png")
            elif elem.name is None and elem.string and elem.string.strip():
                doc.add_paragraph(elem.string.strip())

        if os.path.exists("temp_image.png"):
            os.remove("temp_image.png")

        doc.save("documento_gerado.docx")
        return FileResponse("documento_gerado.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="documento_gerado.docx")

    elif formato == "txt":
        texto = soup.get_text(separator="\n", strip=True)
        with open("documento_gerado.txt", "w", encoding="utf-8") as f:
            f.write(texto)
        return FileResponse("documento_gerado.txt", media_type="text/plain", filename="documento_gerado.txt")

    elif formato == "md":
        texto = soup.get_text(separator="\n", strip=True)
        with open("documento_gerado.md", "w", encoding="utf-8") as f:
            f.write(f"# Documento Gerado\n\n{texto}")
        return FileResponse("documento_gerado.md", media_type="text/markdown", filename="documento_gerado.md")

    else:
        return {"erro": "Formato inválido."}
