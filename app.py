import streamlit as st
from fpdf import FPDF
from docx import Document
import pandas as pd
import io

def gerar_pdf(respostas):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Resumo da Entrevista", ln=True, align='C')
    pdf.ln(10)
    for pergunta, resposta in respostas.items():
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, txt=f"{pergunta}:", ln=True)
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, txt=resposta)
        pdf.ln(5)

    # Salva em arquivo temporário
    pdf.output("resumo_temp.pdf")
    with open("resumo_temp.pdf", "rb") as f:
        pdf_bytes = f.read()
    return pdf_bytes


def gerar_word(respostas):
    doc = Document()
    doc.add_heading('Resumo da Entrevista', 0)
    for pergunta, resposta in respostas.items():
        doc.add_heading(pergunta, level=2)
        doc.add_paragraph(resposta)
    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer

def gerar_txt(respostas):
    content = "Resumo da Entrevista\n\n"
    for pergunta, resposta in respostas.items():
        content += f"{pergunta}\n{resposta}\n\n"
    return content.encode('utf-8')

def gerar_readme(respostas):
    content = "# Resumo da Entrevista\n\n"
    for pergunta, resposta in respostas.items():
        content += f"## {pergunta}\n\n{resposta}\n\n"
    return content.encode('utf-8')

def gerar_csv(respostas):
    df = pd.DataFrame(list(respostas.items()), columns=["Pergunta", "Resposta"])
    csv_buffer = io.StringIO()
    df.to_csv(csv_buffer, index=False)
    return csv_buffer.getvalue().encode('utf-8')

# Perguntas padrão
perguntas_padrao = {
    'objetivo_principal': "Qual é o objetivo principal deste projeto?",
    'objetivos_secundarios': "Existem objetivos secundários? Se sim, quais?",
    'resultados_esperados': "Que entregáveis você espera (relatórios, dashboards, insights)?",
    'criterios_sucesso': "Como saberemos que o projeto foi bem-sucedido?",
    'restricoes': "Existem limitações de prazo, orçamento ou recursos?",
    'publico_alvo': "Quem vai usar os resultados da análise? Qual o nível de familiaridade com dados?",
    'descricao_problema': "Que problema estamos tentando resolver?",
    'contexto_problema': "Como esse problema foi identificado e há quanto tempo ele existe?",
    'impacto_problema': "Qual é o impacto desse problema no negócio/processos/resultados?",
    'problemas_relacionados': "Este problema está ligado a outras questões conhecidas?",
    'perguntas_analise': "Que perguntas específicas precisam ser respondidas para resolver o problema?",
    'metricas_kpis': "Existem KPIs ou métricas específicas a acompanhar?",
    'fontes_dados': "De onde virão os dados necessários (bancos de dados, planilhas, sistemas)?",
    'formato_resultados': "Como as respostas devem ser apresentadas (relatório, gráfico, tabela)?"
}

st.title("📜 Sistema de Entrevista de Dados")
st.write("Preencha o formulário abaixo para gerar o resumo da entrevista no formato que desejar. Você pode editar as perguntas se quiser.")

respostas = {}

with st.form("entrevista_form"):
    st.subheader("Entendimento das Expectativas")
    for key in list(perguntas_padrao.keys())[:6]:
        pergunta_editada = st.text_input(f"Editar pergunta:", value=perguntas_padrao[key], key=f"edit_{key}")
        respostas[pergunta_editada] = st.text_input(f"Resposta:", key=f"resposta_{key}")

    st.subheader("Definição do Problema")
    for key in list(perguntas_padrao.keys())[6:10]:
        pergunta_editada = st.text_input(f"Editar pergunta:", value=perguntas_padrao[key], key=f"edit_{key}")
        respostas[pergunta_editada] = st.text_area(f"Resposta:", key=f"resposta_{key}")

    st.subheader("Definição das Perguntas de Análise")
    for key in list(perguntas_padrao.keys())[10:]:
        pergunta_editada = st.text_input(f"Editar pergunta:", value=perguntas_padrao[key], key=f"edit_{key}")
        respostas[pergunta_editada] = st.text_input(f"Resposta:", key=f"resposta_{key}")

    submitted = st.form_submit_button("Gerar opções de download")

if submitted:
    st.success("Escolha o formato para gerar o arquivo:")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.download_button(
            "📥 Baixar PDF",
            data=gerar_pdf(respostas),
            file_name="resumo_entrevista.pdf",
            mime="application/pdf"
        )

    with col2:
        st.download_button(
            "📄 Baixar Word",
            data=gerar_word(respostas),
            file_name="resumo_entrevista.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with col3:
        st.download_button(
            "📄 Baixar TXT",
            data=gerar_txt(respostas),
            file_name="resumo_entrevista.txt",
            mime="text/plain"
        )

    col4, col5 = st.columns(2)

    with col4:
        st.download_button(
            "📄 Baixar README",
            data=gerar_readme(respostas),
            file_name="README.md",
            mime="text/markdown"
        )

    with col5:
        st.download_button(
            "📄 Baixar CSV",
            data=gerar_csv(respostas),
            file_name="resumo_entrevista.csv",
            mime="text/csv"
        )
